"""Parse FAQ / product documents (PDF, DOCX, TXT, MD) into retrieval chunks.

Only the text layer is read: text-based PDFs work out of the box, scanned
PDFs would need an external OCR step before upload.
"""

from __future__ import annotations

import logging
import os
import re
from typing import List

logger = logging.getLogger("supportops.doc_parser")

SUPPORTED_EXTENSIONS = (".pdf", ".docx", ".txt", ".md", ".markdown")

_MD_IMAGE_RE = re.compile(r"!\[([^\]]*)\]\([^)]*\)")
_MD_LINK_RE = re.compile(r"\[([^\]]+)\]\([^)]*\)")
_MD_MARKUP_RE = re.compile(r"^[ \t]*(#{1,6}|[-*+]|>+|\d+\.)[ \t]+", re.MULTILINE)
_MD_FENCE_RE = re.compile(r"^```.*$", re.MULTILINE)
_SENTENCE_SPLIT_RE = re.compile(r"(?<=[。！？!?；;])|(?<=[.!?] )")


def is_supported_file(file_name: str) -> bool:
    return file_name.lower().endswith(SUPPORTED_EXTENSIONS)


def extract_text(file_path: str) -> str:
    """Return the plain text of a supported document, '' when nothing found."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return _extract_pdf(file_path)
    if ext == ".docx":
        return _extract_docx(file_path)
    if ext in (".md", ".markdown"):
        return _strip_markdown(_read_text(file_path))
    if ext == ".txt":
        return _read_text(file_path)
    raise ValueError(f"不支持的文档类型: {ext or file_path}（支持 PDF/DOCX/TXT/MD）")


def chunk_text(text: str, max_chars: int = 600, overlap: int = 80) -> List[str]:
    """Split text into paragraph-aligned chunks of roughly ``max_chars``.

    Long paragraphs are split at sentence boundaries; consecutive short
    paragraphs are merged. ``overlap`` characters from the end of a chunk are
    carried into the next one to preserve context across boundaries.
    """
    normalized = re.sub(r"[ \t]+", " ", text or "").strip()
    if not normalized:
        return []

    pieces: List[str] = []
    for paragraph in re.split(r"\n\s*\n|\n", normalized):
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        if len(paragraph) <= max_chars:
            pieces.append(paragraph)
        else:
            pieces.extend(_split_long_paragraph(paragraph, max_chars))

    chunks: List[str] = []
    current = ""
    for piece in pieces:
        if not current:
            current = piece
        elif len(current) + len(piece) + 1 <= max_chars:
            current = f"{current}\n{piece}"
        else:
            chunks.append(current)
            tail = current[-overlap:] if overlap > 0 else ""
            current = f"{tail}{piece}" if tail and len(tail) + len(piece) <= max_chars else piece
    if current:
        chunks.append(current)
    return chunks


def _split_long_paragraph(paragraph: str, max_chars: int) -> List[str]:
    sentences = [s for s in _SENTENCE_SPLIT_RE.split(paragraph) if s and s.strip()]
    pieces: List[str] = []
    current = ""
    for sentence in sentences:
        if len(current) + len(sentence) <= max_chars:
            current += sentence
            continue
        if current:
            pieces.append(current.strip())
        while len(sentence) > max_chars:  # pathological sentence, hard split
            pieces.append(sentence[:max_chars])
            sentence = sentence[max_chars:]
        current = sentence
    if current.strip():
        pieces.append(current.strip())
    return pieces


def _extract_pdf(file_path: str) -> str:
    import pdfplumber

    pages: List[str] = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            try:
                pages.append(page.extract_text() or "")
            except Exception as exc:
                logger.warning("PDF page %s parse failed: %s", page.page_number, exc)
    return "\n\n".join(filter(None, pages))


def _extract_docx(file_path: str) -> str:
    from docx import Document

    document = Document(file_path)
    parts: List[str] = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _read_text(file_path: str) -> str:
    with open(file_path, "rb") as handle:
        content = handle.read()
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def _strip_markdown(text: str) -> str:
    text = _MD_FENCE_RE.sub("", text)
    text = _MD_IMAGE_RE.sub(r"\1", text)
    text = _MD_LINK_RE.sub(r"\1", text)
    text = _MD_MARKUP_RE.sub("", text)
    return text
